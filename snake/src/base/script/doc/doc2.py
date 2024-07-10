from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt


def set_cell_vertical_alignment(cell, align="center"):
    """ 设置单元格的垂直居中 """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val = align.upper()  # 将对齐方式转换为大写
    tcValign = OxmlElement('w:vAlign')
    tcValign.set(qn('w:val'), val)
    tcPr.append(tcValign)


def create_table(document, data, columns):
    # 确定行数和列数
    rows = len(data)
    cols = max(sum(cell.get('size', 1) for cell in row) for row in data)

    # 创建表格
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 将整个表格居中

    # 填充表格内容并设置单元格样式
    for i, row in enumerate(data):
        col_index = 0
        for cell_data in row:
            cell = table.cell(i, col_index)
            set_cell_vertical_alignment(cell, align="center")  # 设置垂直居中

            # 设置单元格内容和样式
            cell.text = cell_data.get('value', '')
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cell_data.get('background', 'FFFFFF')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

            paragraph = cell.paragraphs[0]
            paragraph.alignment = 1 if cell_data.get('alignment', 0) == 1 else 0

            for run in paragraph.runs:
                run.bold = cell_data.get('bold', False)

            # 调整单元格大小
            size = cell_data.get('size', 1)
            if size > 1:
                merged_cell = table.cell(i, col_index + size - 1)
                cell.merge(merged_cell)
                col_index += size
            else:
                col_index += 1

    # 合并单元格
    for col in columns:
        i = 0
        while i < rows - 1:
            if data[i][col]['value'] == data[i + 1][col]['value']:
                a = table.cell(i, col)
                b = table.cell(i + 1, col)
                a.merge(b)
                i += 2
            else:
                i += 1

    # 设置表格上下间距
    for paragraph in document.paragraphs:
        paragraph.space_before = Pt(20)
        paragraph.space_after = Pt(20)


# 使用示例
document = Document()
data = [
    [{"value": "配置规范检查", "background": "#FBE4D5", "alignment": 1, "size": 2, "bold": True}],
    [{"value": "其他数据", "background": "#FFFFFF", "alignment": 1, "size": 1, "bold": False}]
]
columns = [0]
create_table(document, data, columns)
document.save('out/example.docx')
